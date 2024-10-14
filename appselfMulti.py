import streamlit as st
import docx
from PyPDF2 import PdfReader
from io import BytesIO
import openai
import warnings

st.markdown("""
    <style>
    .stApp {
        background-color: #f0f0f5;
        color: #333;
    }
    .stButton>button {
        background-color: #0047ab;
        color: white;
        font-size: 16px;
        border-radius: 10px;
    }
    </style>
    """, unsafe_allow_html=True)

def read_docx(file):
    doc = docx.Document(file)
    full_text = [para.text for para in doc.paragraphs]
    return '\n'.join(full_text)

# Function to read PDF files
def read_pdf(file):
    pdf_reader = PdfReader(file)
    full_text = [page.extract_text() for page in pdf_reader.pages if page.extract_text() is not None]
    return '\n'.join(full_text)

# Function to call OpenAI GPT-3.5 API
def generate_prompt(prompt, content):
    combined_prompt = f"{prompt}\n\nContent:\n{content}"
    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",  # Specify the model to use (can use other models)
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": combined_prompt}
        ],
        max_tokens=200
    )
    return response.choices[0].message["content"].strip()

def key_information(content):
    prompt = f"Extract the following details from the legal document: \
          1. Case Number \
          2. Parties Involved \
          3. Date of Filing \
          4. Court Name \
          5. Judgment Dates. \
          Here is the document: {content}"
    
    response = openai.Completion.create(
        model = "gpt-4o-mini",
        messages = [
            {
                "role" : "system",
                "content" : "You are a highly skilled Legal Document Reader, Your role is to extract Key Information from the Document"
            },
            {
                "role" : "user",
                "content" : prompt
            }
        ],
        max_tokens = 200
    )
    return response.choices[0].message["content"].strip()

# Summary Generation Function
def generate_summary(content, size):
    systemMessage = {
                "role": "system",
                "content": (
                    "You are a highly skilled legal document summarizer. "
                    "Your role is to summarize lengthy legal documents while maintaining the integrity and accuracy of the legal content. "
                    "The user will request either a short or long summary. "
                    "- For short summaries, focus on key points and conclusions, writing them concisely. "
                    "- For long summaries, provide a detailed overview of all critical aspects, including main arguments, rulings, and legal precedents. "
                    "Maintain a formal tone and use legal terminology where applicable, making the summary useful for professionals in the legal field."
                )
            }
    
    if size == "Short":
        userMessage = {
            "role": "user",
            "content": (
                f"Summarize this legal document briefly, focusing only on the main points and conclusions:\n\n{content}"
            )
        }
    else:
        userMessage = {
            "role": "user",
            "content": (
                f"Provide a detailed summary of this legal document, including key arguments, legal precedents, and significant rulings:\n\n{content}"
            )
        }

    response = openai.ChatCompletion.create(
        model = "gpt-4o-mini",
        messages = [systemMessage, userMessage],
        max_tokens = 350 if size == 'Short' else 600
    )
    return response.choices[0].message["content"].strip()

def entity_recognition(content):

    sysmessage = {
        "role" : "system", 
        "content" : "You are an expert at reading and understanding Legal Documents, Your job is to identify all entities from a legal document namely, Names of people, companies and Law firms in a legal document"
    }

    usermessage = {
        "role" : "user",
        "content" : f"List all the entities and a small line about what the entity is from the following document : {content}"
    }

    response = openai.Completion.create(
        model = "gpt-40-mini",
        messages = [sysmessage, usermessage],
        max_tokens = 200
    )

    return response.choices[0].message["content"].strip()

def citations_precedents(content):
    prompt = f"Extract all legal citations and case precedents mentioned in the following legal document: {content}"

    response = openai.Completion.create(
            engine="gpt-40-mini",
            prompt=prompt,
            max_tokens=300
        )
    
    return response.choices[0].message["content"].strip()

def obligationTracker(content):
    # Feature: Extract and list all the obligations or responsibilities for each party in the contract.
    return

# These require multi-document uploading

def multiDocumentQnA(content):
    # Just add all documents as a part of prompt 
    return

def changeDetection(content):
    # Feature: Track changes made to contracts over time, detect new versions, and highlight differences. 
    # How it Helps: Lawyers need to ensure that the terms haven't been altered unfavorably, and having automatic change detection 
    # makes contract management easier
    return

# Set your OpenAI API key
openai.api_key = "sk-j7dHvE4yL0ctZihKduJaMxiEVDoORtYp-9F6j7cqTkT3BlbkFJPg19jl3Cdhs5jN8638lCoUw84PoKN5f0SRjlONIsYA"

# File uploader for PDF and DOCX files
uploaded_files = st.file_uploader("Upload a file (PDF or DOCX)", type=["pdf", "docx"], accept_multiple_files=True)

# Function to read DOCX files
def read_docx(file):
    doc = docx.Document(file)
    full_text = [para.text for para in doc.paragraphs]
    return '\n'.join(full_text)

# Function to read PDF files
def read_pdf(file):
    pdf_reader = PdfReader(file)
    full_text = [page.extract_text() for page in pdf_reader.pages if page.extract_text() is not None]
    return '\n'.join(full_text)

documents = {}

if uploaded_files:
    for file in uploaded_files:
        st.write(f"**Filename:** {file.name}") # Of Type bytes
        if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            file_content = read_docx(BytesIO(file.read()))
        elif file.type == "application/pdf":
            file_content = read_pdf(BytesIO(file.read()))
        else:
            file_content = None
        if file_content:
            documents[file.name] = file_content

if documents:
    selected = st.selectbox("Select a document to view", documents.keys())

    if selected:
        selected_content = documents[selected]
        text_area_height = min(400, max(100, len(selected_content.split('\n')) * 20))
        st.text_area("File Content", file_content, height=text_area_height)

    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(["Summary", "QnA", "Key Information", "Entity Recognition", "Citations and Precedents", "Change Detection", "Multi-Document QnA"])

    with tab1:
        st.header("Summarizer")
        st.write("Select the size of the summary and click the Generate Summary button to generate the summary: ")

        selected = st.selectbox("Select a document to view", documents.keys())

        if selected:
            
            file_content = documents[selected]
            col1, col2 = st.columns(2)     
            with col1:
                summsize = st.selectbox("Size of the summary: ", ['Short', 'Long'])

            with col2:
                with st.expander("Current Parameter"):
                    st.write("Size of summary : ", summsize)

            if st.button("Generate Summary : "):
                st.text_area("Summary : ", generate_summary(file_content, summsize), height = 300) 

    with tab2:
        st.header("Single Document, Question and Answer")
        st.write("Document Qna")

        selected = st.selectbox("Select a document to view", documents.keys())

        if selected:

            st.header("Question and Answer")
            st.write("Write a Prompt to generate content based on the Document Provided : ")

            prompt = st.text_area("Enter your prompt related to the file content:")

            file_content = documents[selected]

            if st.button("Create content"):
                if prompt:
                    with st.spinner('Generating response...'):
                        gpt_response = generate_prompt(prompt, file_content)
                        st.text_area("GPT-4o-mini Response:", gpt_response, height=300)
                else:
                    st.write("Please enter a prompt.")
        
    with tab3:

        st.header("Key Information")
        st.write("Generate Key information about the document: Case Number, Parties Involved, Dates and Court Name")

        selected = st.selectbox("Select a document to view", documents.keys())

        if selected:
            file_content = documents[selected]

            if st.button("Get Key Information"):
                gpt_response = key_information(file_content)

                st.text_area("GPT-4o-mini Response:", gpt_response, height=300)
    
    with tab4:

        st.header("Entity Recognition")
        st.write("Identify all the different Entities in the following Document : ")

        selected = st.selectbox("Select a document to view", documents.keys())

        if selected:

            file_content = documents[selected]

            if st.button("Recognize Entities"):
                gpt_response = entity_recognition(file_content)

                st.text_area("GPT-4o-mini Response:", gpt_response, height=300)
    
    with tab5:
        st.header("Citations Precedents")
        st.write("Extract and list all legal citations or case precedents mentioned in the document.")

        selected = st.selectbox("Select a document to view", documents.keys())

        if selected:
            file_content = documents[selected]

            if st.button("Get Citations and Precedents"):
                gpt_response = citations_precedents(file_content)

                st.text_area("GPT-4o-mini Response:", gpt_response, height=300)

    with tab6:
        st.header("Change Detection")
        st.write("Detect Version Changes over multiple documents : ")

        selected = st.multiselect("Select documents", documents.keys(), max_selections=2)

        docs = []
        if selected:

            for doc in selected:
                docs.append(documents[selected])

            if st.button("Get Changes in file"):
                gpt_response = changeDetection(file_content)

                st.text_area("GPT-4o-mini Response:", gpt_response, height=300)
    
    with tab7:

        st.header("Multi Document QnA : ")
        st.write("Add multiple documents for question and answers")

        selected = st.multiselect("Select Documents : ", documents.keys())

        docs = []
        if selected:

            for doc in selected:
                docs.append(documents[selected])
            
            if st.button("Ask questions to multiple documents : "):

                if st.button("Create content"):
                    if prompt:
                        with st.spinner('Generating response...'):
                            gpt_response = multiDocumentQnA(prompt, file_content)
                            st.text_area("GPT-4o-mini Response:", gpt_response, height=300)
                    else:
                        st.write("Please enter a prompt.")
else:
    st.write("Please upload a file.")
    